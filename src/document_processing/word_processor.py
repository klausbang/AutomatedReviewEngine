"""
MS Word Document Processor

This module handles MS Word document reading, text extraction, and structure analysis
for the Automated Review Engine.
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from docx import Document
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from io import BytesIO

# Configure logging
logger = logging.getLogger(__name__)

@dataclass
class WordSection:
    """Represents a section found in a Word document"""
    title: str
    content: str
    level: int  # heading level (1-9)
    style: str  # paragraph style name
    position: int  # position in document

@dataclass
class WordTable:
    """Represents a table found in a Word document"""
    position: int
    headers: List[str]
    rows: List[List[str]]
    style: Optional[str] = None

@dataclass
class WordProcessingResult:
    """Result of Word document processing"""
    success: bool
    content: str
    sections: List[WordSection]
    tables: List[WordTable]
    metadata: Dict[str, Any]
    paragraph_count: int
    error_message: Optional[str] = None

class WordProcessor:
    """
    MS Word document processor for extracting text and structure from DOCX files.
    
    Uses python-docx library for comprehensive Word document processing:
    - Text extraction from paragraphs
    - Table extraction and parsing
    - Heading and section detection
    - Document metadata extraction
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize Word processor with configuration.
        
        Args:
            config: Configuration dictionary with processing options
        """
        self.config = config or {}
        self.max_file_size_mb = self.config.get('max_file_size_mb', 50)
        self.extract_tables = self.config.get('extract_tables', True)
        self.extract_sections = self.config.get('extract_sections', True)
        self.preserve_formatting = self.config.get('preserve_formatting', False)
        
    def process_word(self, file_path: Path) -> WordProcessingResult:
        """
        Process a Word document and extract content, structure, and metadata.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            WordProcessingResult with extracted content and structure
        """
        try:
            # Validate file
            validation_result = self._validate_file(file_path)
            if not validation_result['valid']:
                return WordProcessingResult(
                    success=False,
                    content="",
                    sections=[],
                    tables=[],
                    metadata={},
                    paragraph_count=0,
                    error_message=validation_result['error']
                )
            
            # Open and process document
            doc = Document(file_path)
            
            # Extract content and structure
            content = self._extract_content(doc)
            sections = self._extract_sections(doc) if self.extract_sections else []
            tables = self._extract_tables(doc) if self.extract_tables else []
            metadata = self._extract_metadata(doc, file_path)
            paragraph_count = len(doc.paragraphs)
            
            return WordProcessingResult(
                success=True,
                content=content,
                sections=sections,
                tables=tables,
                metadata=metadata,
                paragraph_count=paragraph_count
            )
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}")
            return WordProcessingResult(
                success=False,
                content="",
                sections=[],
                tables=[],
                metadata={},
                paragraph_count=0,
                error_message=f"Processing error: {str(e)}"
            )
    
    def process_word_from_bytes(self, word_bytes: bytes, filename: str = "uploaded.docx") -> WordProcessingResult:
        """
        Process Word document from bytes (useful for uploaded files).
        
        Args:
            word_bytes: Word file as bytes
            filename: Original filename for reference
            
        Returns:
            WordProcessingResult with extracted content and structure
        """
        try:
            # Validate file size
            file_size_mb = len(word_bytes) / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                return WordProcessingResult(
                    success=False,
                    content="",
                    sections=[],
                    tables=[],
                    metadata={},
                    paragraph_count=0,
                    error_message=f"File too large: {file_size_mb:.1f}MB exceeds limit of {self.max_file_size_mb}MB"
                )
            
            # Open document from bytes
            doc_file = BytesIO(word_bytes)
            doc = Document(doc_file)
            
            # Extract content and structure
            content = self._extract_content(doc)
            sections = self._extract_sections(doc) if self.extract_sections else []
            tables = self._extract_tables(doc) if self.extract_tables else []
            metadata = self._extract_metadata_from_bytes(doc, filename)
            paragraph_count = len(doc.paragraphs)
            
            return WordProcessingResult(
                success=True,
                content=content,
                sections=sections,
                tables=tables,
                metadata=metadata,
                paragraph_count=paragraph_count
            )
            
        except Exception as e:
            logger.error(f"Error processing Word bytes for {filename}: {str(e)}")
            return WordProcessingResult(
                success=False,
                content="",
                sections=[],
                tables=[],
                metadata={},
                paragraph_count=0,
                error_message=f"Processing error: {str(e)}"
            )
    
    def _validate_file(self, file_path: Path) -> Dict[str, Any]:
        """Validate Word file before processing"""
        try:
            if not file_path.exists():
                return {'valid': False, 'error': 'File does not exist'}
            
            # Check file size
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                return {'valid': False, 'error': f'File too large: {file_size_mb:.1f}MB exceeds limit of {self.max_file_size_mb}MB'}
            
            # Check file extension
            if file_path.suffix.lower() not in ['.docx', '.doc']:
                return {'valid': False, 'error': 'File is not a Word document'}
            
            # Note: .doc files are not supported by python-docx, only .docx
            if file_path.suffix.lower() == '.doc':
                return {'valid': False, 'error': 'Legacy .doc format not supported. Please use .docx format.'}
            
            return {'valid': True}
            
        except Exception as e:
            return {'valid': False, 'error': f'Validation error: {str(e)}'}
    
    def _extract_content(self, doc: Document) -> str:
        """Extract all text content from the document"""
        content_parts = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                # Find corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            if self.preserve_formatting:
                                # Add style information
                                style_name = para.style.name if para.style else "Normal"
                                content_parts.append(f"[{style_name}] {text}")
                            else:
                                content_parts.append(text)
                        break
            
            elif element.tag.endswith('tbl'):  # Table
                # Find corresponding table object
                for table in doc.tables:
                    if table._element == element:
                        if self.extract_tables:
                            content_parts.append(self._table_to_text(table))
                        break
        
        return '\n\n'.join(content_parts)
    
    def _extract_sections(self, doc: Document) -> List[WordSection]:
        """Extract sections based on heading styles"""
        sections = []
        current_section = None
        current_content = []
        
        for position, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else "Normal"
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if this is a heading
            heading_level = self._get_heading_level(style_name)
            
            if heading_level > 0:
                # Save previous section
                if current_section:
                    section = WordSection(
                        title=current_section['title'],
                        content='\n'.join(current_content),
                        level=current_section['level'],
                        style=current_section['style'],
                        position=current_section['position']
                    )
                    sections.append(section)
                
                # Start new section
                current_section = {
                    'title': text,
                    'level': heading_level,
                    'style': style_name,
                    'position': position
                }
                current_content = []
            else:
                # Add to current section content
                if current_section:
                    current_content.append(text)
                else:
                    # Content before first heading
                    if not sections:  # Create a default section for initial content
                        current_section = {
                            'title': 'Document Content',
                            'level': 1,
                            'style': 'Normal',
                            'position': 0
                        }
                        current_content = [text]
        
        # Add last section
        if current_section:
            section = WordSection(
                title=current_section['title'],
                content='\n'.join(current_content),
                level=current_section['level'],
                style=current_section['style'],
                position=current_section['position']
            )
            sections.append(section)
        
        return sections
    
    def _extract_tables(self, doc: Document) -> List[WordTable]:
        """Extract all tables from the document"""
        tables = []
        
        for position, table in enumerate(doc.tables):
            try:
                # Extract headers (first row)
                headers = []
                if len(table.rows) > 0:
                    header_row = table.rows[0]
                    headers = [cell.text.strip() for cell in header_row.cells]
                
                # Extract data rows
                rows = []
                for row in table.rows[1:]:  # Skip header row
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows.append(row_data)
                
                # Get table style if available
                table_style = None
                try:
                    table_style = table.style.name if table.style else None
                except:
                    pass
                
                word_table = WordTable(
                    position=position,
                    headers=headers,
                    rows=rows,
                    style=table_style
                )
                tables.append(word_table)
                
            except Exception as e:
                logger.warning(f"Could not process table {position}: {str(e)}")
                continue
        
        return tables
    
    def _extract_metadata(self, doc: Document, file_path: Path) -> Dict[str, Any]:
        """Extract document metadata"""
        metadata = {
            'filename': file_path.name,
            'file_size': file_path.stat().st_size,
            'file_path': str(file_path)
        }
        
        # Extract core properties
        core_props = doc.core_properties
        if core_props:
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'category': core_props.category or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'version': core_props.version or '',
                'revision': core_props.revision or ''
            })
        
        # Document statistics
        metadata.update({
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'section_count': len(doc.sections)
        })
        
        return metadata
    
    def _extract_metadata_from_bytes(self, doc: Document, filename: str) -> Dict[str, Any]:
        """Extract document metadata from bytes"""
        metadata = {
            'filename': filename,
            'file_size': 0,  # Not available from bytes
            'file_path': 'uploaded'
        }
        
        # Extract core properties
        core_props = doc.core_properties
        if core_props:
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'category': core_props.category or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'version': core_props.version or '',
                'revision': core_props.revision or ''
            })
        
        # Document statistics
        metadata.update({
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'section_count': len(doc.sections)
        })
        
        return metadata
    
    def _get_heading_level(self, style_name: str) -> int:
        """Determine heading level from style name"""
        if not style_name:
            return 0
        
        # Standard heading styles
        heading_styles = {
            'Heading 1': 1, 'Title': 1,
            'Heading 2': 2, 'Subtitle': 2,
            'Heading 3': 3,
            'Heading 4': 4,
            'Heading 5': 5,
            'Heading 6': 6,
            'Heading 7': 7,
            'Heading 8': 8,
            'Heading 9': 9
        }
        
        return heading_styles.get(style_name, 0)
    
    def _table_to_text(self, table: DocxTable) -> str:
        """Convert table to text representation"""
        text_parts = []
        
        try:
            # Add table header
            text_parts.append("--- TABLE ---")
            
            for row_num, row in enumerate(table.rows):
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):  # Only add non-empty rows
                    if row_num == 0:
                        # Header row
                        text_parts.append("Headers: " + " | ".join(row_cells))
                    else:
                        # Data row
                        text_parts.append(f"Row {row_num}: " + " | ".join(row_cells))
            
            text_parts.append("--- END TABLE ---")
            
        except Exception as e:
            logger.warning(f"Could not convert table to text: {str(e)}")
            text_parts.append("--- TABLE (could not parse) ---")
        
        return '\n'.join(text_parts)
    
    def get_document_statistics(self, doc: Document) -> Dict[str, int]:
        """Get comprehensive document statistics"""
        stats = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'sections': len(doc.sections),
            'headings': 0,
            'words': 0,
            'characters': 0
        }
        
        # Count headings and words
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Count words and characters
                words = len(text.split())
                stats['words'] += words
                stats['characters'] += len(text)
                
                # Check if heading
                style_name = para.style.name if para.style else "Normal"
                if self._get_heading_level(style_name) > 0:
                    stats['headings'] += 1
        
        return stats
