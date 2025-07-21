"""
Document Validator

This module provides validation utilities for documents processed by the
Automated Review Engine.
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass
from enum import Enum
import magic
import mimetypes

# Configure logging
logger = logging.getLogger(__name__)

class DocumentType(Enum):
    """Supported document types"""
    PDF = "pdf"
    WORD = "docx"
    UNKNOWN = "unknown"

class ValidationSeverity(Enum):
    """Validation issue severity levels"""
    ERROR = "error"
    WARNING = "warning"
    INFO = "info"

@dataclass
class ValidationIssue:
    """Represents a validation issue"""
    severity: ValidationSeverity
    code: str
    message: str
    details: Optional[str] = None

@dataclass
class ValidationResult:
    """Result of document validation"""
    valid: bool
    document_type: DocumentType
    file_size_mb: float
    issues: List[ValidationIssue]
    metadata: Dict[str, Any]

class DocumentValidator:
    """
    Comprehensive document validator for the Automated Review Engine.
    
    Validates document format, size, content, and structure before processing.
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize document validator with configuration.
        
        Args:
            config: Configuration dictionary with validation options
        """
        self.config = config or {}
        self.max_file_size_mb = self.config.get('max_file_size_mb', 50)
        self.supported_formats = self.config.get('supported_formats', ['pdf', 'docx'])
        self.strict_mode = self.config.get('strict_mode', False)
        
        # MIME type mappings
        self.mime_mappings = {
            'application/pdf': DocumentType.PDF,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentType.WORD,
            'application/msword': DocumentType.WORD  # Legacy .doc format
        }
    
    def validate_file(self, file_path: Path) -> ValidationResult:
        """
        Validate a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ValidationResult with validation status and issues
        """
        issues = []
        metadata = {}
        
        try:
            # Basic file existence check
            if not file_path.exists():
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="FILE_NOT_FOUND",
                    message="File does not exist",
                    details=f"Path: {file_path}"
                ))
                return ValidationResult(
                    valid=False,
                    document_type=DocumentType.UNKNOWN,
                    file_size_mb=0.0,
                    issues=issues,
                    metadata=metadata
                )
            
            # File size validation
            file_size_bytes = file_path.stat().st_size
            file_size_mb = file_size_bytes / (1024 * 1024)
            metadata['file_size_bytes'] = file_size_bytes
            metadata['file_size_mb'] = file_size_mb
            
            if file_size_mb > self.max_file_size_mb:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="FILE_TOO_LARGE",
                    message=f"File size {file_size_mb:.1f}MB exceeds limit of {self.max_file_size_mb}MB",
                    details=f"File size: {file_size_mb:.2f}MB"
                ))
            
            if file_size_mb == 0:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="EMPTY_FILE",
                    message="File is empty",
                    details="File size is 0 bytes"
                ))
            
            # Document type detection
            document_type = self._detect_document_type(file_path)
            metadata['detected_type'] = document_type.value
            
            # Format validation
            if document_type == DocumentType.UNKNOWN:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="UNSUPPORTED_FORMAT",
                    message="Document format not supported",
                    details=f"Supported formats: {', '.join(self.supported_formats)}"
                ))
            elif document_type.value not in self.supported_formats:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="FORMAT_NOT_ENABLED",
                    message=f"Format {document_type.value} not enabled in configuration",
                    details=f"Enabled formats: {', '.join(self.supported_formats)}"
                ))
            
            # File extension validation
            extension_issues = self._validate_file_extension(file_path, document_type)
            issues.extend(extension_issues)
            
            # Content validation (basic)
            content_issues = self._validate_file_content(file_path, document_type)
            issues.extend(content_issues)
            
            # Security validation
            security_issues = self._validate_file_security(file_path)
            issues.extend(security_issues)
            
            # Determine overall validity
            has_errors = any(issue.severity == ValidationSeverity.ERROR for issue in issues)
            valid = not has_errors
            
            return ValidationResult(
                valid=valid,
                document_type=document_type,
                file_size_mb=file_size_mb,
                issues=issues,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Validation error for {file_path}: {str(e)}")
            issues.append(ValidationIssue(
                severity=ValidationSeverity.ERROR,
                code="VALIDATION_ERROR",
                message="Unexpected validation error",
                details=str(e)
            ))
            
            return ValidationResult(
                valid=False,
                document_type=DocumentType.UNKNOWN,
                file_size_mb=0.0,
                issues=issues,
                metadata=metadata
            )
    
    def validate_bytes(self, file_bytes: bytes, filename: str) -> ValidationResult:
        """
        Validate document from bytes.
        
        Args:
            file_bytes: Document content as bytes
            filename: Original filename for reference
            
        Returns:
            ValidationResult with validation status and issues
        """
        issues = []
        metadata = {'filename': filename}
        
        try:
            # File size validation
            file_size_bytes = len(file_bytes)
            file_size_mb = file_size_bytes / (1024 * 1024)
            metadata['file_size_bytes'] = file_size_bytes
            metadata['file_size_mb'] = file_size_mb
            
            if file_size_mb > self.max_file_size_mb:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="FILE_TOO_LARGE",
                    message=f"File size {file_size_mb:.1f}MB exceeds limit of {self.max_file_size_mb}MB",
                    details=f"File size: {file_size_mb:.2f}MB"
                ))
            
            if file_size_mb == 0:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="EMPTY_FILE",
                    message="File is empty",
                    details="File size is 0 bytes"
                ))
            
            # Document type detection from bytes
            document_type = self._detect_document_type_from_bytes(file_bytes, filename)
            metadata['detected_type'] = document_type.value
            
            # Format validation
            if document_type == DocumentType.UNKNOWN:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="UNSUPPORTED_FORMAT",
                    message="Document format not supported",
                    details=f"Supported formats: {', '.join(self.supported_formats)}"
                ))
            elif document_type.value not in self.supported_formats:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="FORMAT_NOT_ENABLED",
                    message=f"Format {document_type.value} not enabled in configuration",
                    details=f"Enabled formats: {', '.join(self.supported_formats)}"
                ))
            
            # Extension validation
            extension_issues = self._validate_filename_extension(filename, document_type)
            issues.extend(extension_issues)
            
            # Content validation from bytes
            content_issues = self._validate_bytes_content(file_bytes, document_type)
            issues.extend(content_issues)
            
            # Determine overall validity
            has_errors = any(issue.severity == ValidationSeverity.ERROR for issue in issues)
            valid = not has_errors
            
            return ValidationResult(
                valid=valid,
                document_type=document_type,
                file_size_mb=file_size_mb,
                issues=issues,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Validation error for {filename}: {str(e)}")
            issues.append(ValidationIssue(
                severity=ValidationSeverity.ERROR,
                code="VALIDATION_ERROR",
                message="Unexpected validation error",
                details=str(e)
            ))
            
            return ValidationResult(
                valid=False,
                document_type=DocumentType.UNKNOWN,
                file_size_mb=0.0,
                issues=issues,
                metadata=metadata
            )
    
    def _detect_document_type(self, file_path: Path) -> DocumentType:
        """Detect document type from file"""
        try:
            # Try using python-magic for MIME type detection
            try:
                mime_type = magic.from_file(str(file_path), mime=True)
                if mime_type in self.mime_mappings:
                    return self.mime_mappings[mime_type]
            except Exception as e:
                logger.debug(f"python-magic detection failed: {e}")
            
            # Fallback to mimetypes module
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if mime_type in self.mime_mappings:
                return self.mime_mappings[mime_type]
            
            # Fallback to file extension
            extension = file_path.suffix.lower()
            extension_mappings = {
                '.pdf': DocumentType.PDF,
                '.docx': DocumentType.WORD,
                '.doc': DocumentType.WORD
            }
            
            return extension_mappings.get(extension, DocumentType.UNKNOWN)
            
        except Exception as e:
            logger.warning(f"Document type detection failed: {e}")
            return DocumentType.UNKNOWN
    
    def _detect_document_type_from_bytes(self, file_bytes: bytes, filename: str) -> DocumentType:
        """Detect document type from bytes"""
        try:
            # Try using python-magic
            try:
                mime_type = magic.from_buffer(file_bytes, mime=True)
                if mime_type in self.mime_mappings:
                    return self.mime_mappings[mime_type]
            except Exception as e:
                logger.debug(f"python-magic detection from bytes failed: {e}")
            
            # Check file signature (magic bytes)
            if file_bytes.startswith(b'%PDF'):
                return DocumentType.PDF
            
            # DOCX files are ZIP archives with specific structure
            if file_bytes.startswith(b'PK\x03\x04') or file_bytes.startswith(b'PK\x05\x06') or file_bytes.startswith(b'PK\x07\x08'):
                # Could be DOCX (or other ZIP-based format)
                # Additional check would be needed to confirm DOCX
                return DocumentType.WORD
            
            # Fallback to filename extension
            path = Path(filename)
            extension = path.suffix.lower()
            extension_mappings = {
                '.pdf': DocumentType.PDF,
                '.docx': DocumentType.WORD,
                '.doc': DocumentType.WORD
            }
            
            return extension_mappings.get(extension, DocumentType.UNKNOWN)
            
        except Exception as e:
            logger.warning(f"Document type detection from bytes failed: {e}")
            return DocumentType.UNKNOWN
    
    def _validate_file_extension(self, file_path: Path, document_type: DocumentType) -> List[ValidationIssue]:
        """Validate file extension matches detected type"""
        issues = []
        extension = file_path.suffix.lower()
        
        expected_extensions = {
            DocumentType.PDF: ['.pdf'],
            DocumentType.WORD: ['.docx']  # Note: .doc not supported by python-docx
        }
        
        if document_type in expected_extensions:
            if extension not in expected_extensions[document_type]:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.WARNING,
                    code="EXTENSION_MISMATCH",
                    message=f"File extension '{extension}' doesn't match detected type '{document_type.value}'",
                    details=f"Expected extensions: {', '.join(expected_extensions[document_type])}"
                ))
        
        # Special handling for .doc files
        if extension == '.doc':
            issues.append(ValidationIssue(
                severity=ValidationSeverity.ERROR,
                code="LEGACY_FORMAT",
                message="Legacy .doc format not supported",
                details="Please convert to .docx format"
            ))
        
        return issues
    
    def _validate_filename_extension(self, filename: str, document_type: DocumentType) -> List[ValidationIssue]:
        """Validate filename extension matches detected type"""
        path = Path(filename)
        return self._validate_file_extension(path, document_type)
    
    def _validate_file_content(self, file_path: Path, document_type: DocumentType) -> List[ValidationIssue]:
        """Validate file content structure"""
        issues = []
        
        try:
            if document_type == DocumentType.PDF:
                issues.extend(self._validate_pdf_content(file_path))
            elif document_type == DocumentType.WORD:
                issues.extend(self._validate_word_content(file_path))
                
        except Exception as e:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.WARNING,
                code="CONTENT_VALIDATION_ERROR",
                message="Could not validate file content",
                details=str(e)
            ))
        
        return issues
    
    def _validate_bytes_content(self, file_bytes: bytes, document_type: DocumentType) -> List[ValidationIssue]:
        """Validate content from bytes"""
        issues = []
        
        try:
            if document_type == DocumentType.PDF:
                issues.extend(self._validate_pdf_bytes(file_bytes))
            elif document_type == DocumentType.WORD:
                issues.extend(self._validate_word_bytes(file_bytes))
                
        except Exception as e:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.WARNING,
                code="CONTENT_VALIDATION_ERROR",
                message="Could not validate file content",
                details=str(e)
            ))
        
        return issues
    
    def _validate_pdf_content(self, file_path: Path) -> List[ValidationIssue]:
        """Validate PDF-specific content"""
        issues = []
        
        try:
            # Basic PDF validation - try to open
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Check if PDF has pages
                if len(reader.pages) == 0:
                    issues.append(ValidationIssue(
                        severity=ValidationSeverity.ERROR,
                        code="EMPTY_PDF",
                        message="PDF has no pages",
                        details="Document contains no content"
                    ))
                
                # Check if PDF is encrypted
                if reader.is_encrypted:
                    issues.append(ValidationIssue(
                        severity=ValidationSeverity.ERROR,
                        code="ENCRYPTED_PDF",
                        message="PDF is encrypted",
                        details="Password-protected documents are not supported"
                    ))
                
        except Exception as e:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.ERROR,
                code="INVALID_PDF",
                message="PDF file appears to be corrupted",
                details=str(e)
            ))
        
        return issues
    
    def _validate_pdf_bytes(self, file_bytes: bytes) -> List[ValidationIssue]:
        """Validate PDF from bytes"""
        issues = []
        
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_file = BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(pdf_file)
            
            # Check if PDF has pages
            if len(reader.pages) == 0:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="EMPTY_PDF",
                    message="PDF has no pages",
                    details="Document contains no content"
                ))
            
            # Check if PDF is encrypted
            if reader.is_encrypted:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.ERROR,
                    code="ENCRYPTED_PDF",
                    message="PDF is encrypted",
                    details="Password-protected documents are not supported"
                ))
                
        except Exception as e:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.ERROR,
                code="INVALID_PDF",
                message="PDF file appears to be corrupted",
                details=str(e)
            ))
        
        return issues
    
    def _validate_word_content(self, file_path: Path) -> List[ValidationIssue]:
        """Validate Word document content"""
        issues = []
        
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Check if document has content
            if len(doc.paragraphs) == 0:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.WARNING,
                    code="EMPTY_DOCUMENT",
                    message="Document has no paragraphs",
                    details="Document may be empty"
                ))
            
            # Check for actual text content
            has_text = any(para.text.strip() for para in doc.paragraphs)
            if not has_text:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.WARNING,
                    code="NO_TEXT_CONTENT",
                    message="Document contains no readable text",
                    details="All paragraphs appear to be empty"
                ))
                
        except Exception as e:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.ERROR,
                code="INVALID_WORD_DOCUMENT",
                message="Word document appears to be corrupted",
                details=str(e)
            ))
        
        return issues
    
    def _validate_word_bytes(self, file_bytes: bytes) -> List[ValidationIssue]:
        """Validate Word document from bytes"""
        issues = []
        
        try:
            from docx import Document
            from io import BytesIO
            
            doc_file = BytesIO(file_bytes)
            doc = Document(doc_file)
            
            # Check if document has content
            if len(doc.paragraphs) == 0:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.WARNING,
                    code="EMPTY_DOCUMENT",
                    message="Document has no paragraphs",
                    details="Document may be empty"
                ))
            
            # Check for actual text content
            has_text = any(para.text.strip() for para in doc.paragraphs)
            if not has_text:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.WARNING,
                    code="NO_TEXT_CONTENT",
                    message="Document contains no readable text",
                    details="All paragraphs appear to be empty"
                ))
                
        except Exception as e:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.ERROR,
                code="INVALID_WORD_DOCUMENT",
                message="Word document appears to be corrupted",
                details=str(e)
            ))
        
        return issues
    
    def _validate_file_security(self, file_path: Path) -> List[ValidationIssue]:
        """Basic security validation"""
        issues = []
        
        # Check for suspicious file names
        filename = file_path.name.lower()
        suspicious_patterns = [
            'script', 'exec', 'cmd', 'bat', 'exe', 'dll', 'scr'
        ]
        
        for pattern in suspicious_patterns:
            if pattern in filename:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.WARNING,
                    code="SUSPICIOUS_FILENAME",
                    message=f"Filename contains suspicious pattern: {pattern}",
                    details="Please verify file is safe"
                ))
                break
        
        return issues
    
    def get_validation_summary(self, result: ValidationResult) -> str:
        """Get human-readable validation summary"""
        if result.valid:
            return f"✅ Valid {result.document_type.value.upper()} document ({result.file_size_mb:.1f}MB)"
        
        error_count = sum(1 for issue in result.issues if issue.severity == ValidationSeverity.ERROR)
        warning_count = sum(1 for issue in result.issues if issue.severity == ValidationSeverity.WARNING)
        
        summary = f"❌ Invalid document: {error_count} errors"
        if warning_count > 0:
            summary += f", {warning_count} warnings"
        
        return summary
