"""
Document Analyzer - Automated Review Engine

Advanced document parsing and analysis for PDF and Word documents.
Extracts text, structure, and metadata for review processing.

Phase 3.2: Review Logic - Document Analysis Component
"""

import sys
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Tuple
import re
from datetime import datetime
from dataclasses import dataclass
from enum import Enum

# Add project paths
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / "src"))

# Document processing imports
try:
    import PyPDF2
    import pdfplumber
    from docx import Document as DocxDocument
    from docx.document import Document as DocxDocumentType
    PDF_AVAILABLE = True
    DOCX_AVAILABLE = True
except ImportError:
    PyPDF2 = None
    pdfplumber = None
    DocxDocument = None
    DocxDocumentType = None
    PDF_AVAILABLE = False
    DOCX_AVAILABLE = False

# Core imports
try:
    from src.core.logging_manager import LoggingManager
    from src.core.error_handler import ErrorHandler
    from src.core.config_manager import ConfigManager
except ImportError:
    LoggingManager = None
    ErrorHandler = None
    ConfigManager = None


class DocumentType(Enum):
    """Supported document types"""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    UNKNOWN = "unknown"


class DocumentStructure(Enum):
    """Document structure elements"""
    HEADER = "header"
    PARAGRAPH = "paragraph"
    TABLE = "table"
    LIST = "list"
    FOOTER = "footer"
    METADATA = "metadata"


@dataclass
class DocumentElement:
    """Represents a parsed document element"""
    type: DocumentStructure
    content: str
    position: int
    metadata: Dict[str, Any]
    confidence: float = 1.0


@dataclass
class DocumentMetadata:
    """Document metadata information"""
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    creator: Optional[str] = None
    creation_date: Optional[datetime] = None
    modification_date: Optional[datetime] = None
    page_count: int = 0
    word_count: int = 0
    character_count: int = 0
    language: Optional[str] = None
    file_size: int = 0


@dataclass
class AnalysisResult:
    """Document analysis result"""
    document_type: DocumentType
    text_content: str
    elements: List[DocumentElement]
    metadata: DocumentMetadata
    structure_analysis: Dict[str, Any]
    extraction_errors: List[str]
    processing_time: float
    success: bool


class DocumentAnalyzer:
    """Advanced document analyzer with PDF and Word support"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize document analyzer
        
        Args:
            config: Configuration dictionary
        """
        self.config = config or self._get_default_config()
        self.logger = None
        self.error_handler = None
        
        # Initialize core components if available
        self._initialize_core_components()
        
        # Analysis statistics
        self.analysis_stats = {
            'documents_analyzed': 0,
            'successful_analyses': 0,
            'failed_analyses': 0,
            'total_processing_time': 0.0
        }
    
    def _get_default_config(self) -> Dict[str, Any]:
        """Get default analyzer configuration"""
        return {
            'max_file_size_mb': 50,
            'extract_metadata': True,
            'extract_images': False,
            'preserve_formatting': True,
            'language_detection': True,
            'structure_analysis': True,
            'content_preprocessing': True,
            'supported_formats': ['.pdf', '.docx', '.doc'],
            'pdf_engine': 'pdfplumber',  # 'PyPDF2' or 'pdfplumber'
            'text_extraction_timeout': 120,  # seconds
            'min_confidence_threshold': 0.7
        }
    
    def _initialize_core_components(self):
        """Initialize core infrastructure components"""
        try:
            if LoggingManager:
                self.logger_manager = LoggingManager({'level': 'INFO'})
                self.logger_manager.initialize()
                self.logger = self.logger_manager.get_logger('review.document_analyzer')
                
            if ErrorHandler:
                self.error_handler = ErrorHandler()
                
            if self.logger:
                self.logger.info("Document analyzer initialized successfully")
                
        except Exception as e:
            if self.logger:
                self.logger.error(f"Failed to initialize core components: {e}")
    
    def analyze_document(self, document_path: Union[str, Path]) -> AnalysisResult:
        """
        Analyze a document and extract content, structure, and metadata
        
        Args:
            document_path: Path to document file
            
        Returns:
            AnalysisResult with extracted information
        """
        start_time = datetime.now()
        document_path = Path(document_path)
        
        try:
            # Validate file exists and format
            if not document_path.exists():
                raise FileNotFoundError(f"Document not found: {document_path}")
            
            # Determine document type
            doc_type = self._detect_document_type(document_path)
            
            if doc_type == DocumentType.UNKNOWN:
                raise ValueError(f"Unsupported document format: {document_path.suffix}")
            
            # Validate file size
            file_size = document_path.stat().st_size
            max_size = self.config['max_file_size_mb'] * 1024 * 1024
            
            if file_size > max_size:
                raise ValueError(f"File too large: {file_size} bytes (max: {max_size})")
            
            # Perform document-specific analysis
            if doc_type == DocumentType.PDF:
                result = self._analyze_pdf(document_path)
            elif doc_type in [DocumentType.DOCX, DocumentType.DOC]:
                result = self._analyze_word(document_path)
            else:
                raise ValueError(f"Unsupported document type: {doc_type}")
            
            # Post-process results
            result = self._post_process_analysis(result, document_path, file_size)
            
            # Update statistics
            processing_time = (datetime.now() - start_time).total_seconds()
            result.processing_time = processing_time
            
            self.analysis_stats['documents_analyzed'] += 1
            self.analysis_stats['total_processing_time'] += processing_time
            
            if result.success:
                self.analysis_stats['successful_analyses'] += 1
            else:
                self.analysis_stats['failed_analyses'] += 1
            
            if self.logger:
                self.logger.info(f"Document analysis completed: {document_path.name} ({processing_time:.2f}s)")
            
            return result
            
        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds()
            
            if self.error_handler:
                error_context = self.error_handler.handle_error(e)
                error_message = error_context.user_message
            else:
                error_message = str(e)
            
            if self.logger:
                self.logger.error(f"Document analysis failed: {document_path.name} - {error_message}")
            
            # Return failed analysis result
            return AnalysisResult(
                document_type=DocumentType.UNKNOWN,
                text_content="",
                elements=[],
                metadata=DocumentMetadata(file_size=document_path.stat().st_size if document_path.exists() else 0),
                structure_analysis={},
                extraction_errors=[error_message],
                processing_time=processing_time,
                success=False
            )
    
    def _detect_document_type(self, document_path: Path) -> DocumentType:
        """Detect document type from file extension and content"""
        extension = document_path.suffix.lower()
        
        if extension == '.pdf':
            return DocumentType.PDF
        elif extension == '.docx':
            return DocumentType.DOCX
        elif extension == '.doc':
            return DocumentType.DOC
        else:
            return DocumentType.UNKNOWN
    
    def _analyze_pdf(self, document_path: Path) -> AnalysisResult:
        """Analyze PDF document"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available. Install PyPDF2 and pdfplumber.")
        
        text_content = ""
        elements = []
        metadata = DocumentMetadata()
        errors = []
        
        try:
            # Use pdfplumber for better text extraction
            if self.config['pdf_engine'] == 'pdfplumber' and pdfplumber:
                with pdfplumber.open(document_path) as pdf:
                    # Extract metadata
                    if pdf.metadata:
                        metadata = self._extract_pdf_metadata_pdfplumber(pdf.metadata)
                    
                    metadata.page_count = len(pdf.pages)
                    
                    # Extract text from all pages
                    page_texts = []
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                page_texts.append(page_text)
                                
                                # Create page element
                                elements.append(DocumentElement(
                                    type=DocumentStructure.PARAGRAPH,
                                    content=page_text,
                                    position=page_num,
                                    metadata={'page': page_num + 1, 'type': 'page_content'}
                                ))
                        except Exception as e:
                            errors.append(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                    
                    text_content = "\n\n".join(page_texts)
            
            # Fallback to PyPDF2
            elif PyPDF2:
                with open(document_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    # Extract metadata
                    if pdf_reader.metadata:
                        metadata = self._extract_pdf_metadata_pypdf2(pdf_reader.metadata)
                    
                    metadata.page_count = len(pdf_reader.pages)
                    
                    # Extract text from all pages
                    page_texts = []
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                page_texts.append(page_text)
                                
                                elements.append(DocumentElement(
                                    type=DocumentStructure.PARAGRAPH,
                                    content=page_text,
                                    position=page_num,
                                    metadata={'page': page_num + 1, 'type': 'page_content'}
                                ))
                        except Exception as e:
                            errors.append(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                    
                    text_content = "\n\n".join(page_texts)
            
            else:
                raise ImportError("No PDF processing library available")
            
            # Analyze document structure
            structure_analysis = self._analyze_text_structure(text_content)
            
            # Update metadata with text statistics
            metadata.word_count = len(text_content.split())
            metadata.character_count = len(text_content)
            
        except Exception as e:
            errors.append(f"PDF analysis error: {str(e)}")
        
        return AnalysisResult(
            document_type=DocumentType.PDF,
            text_content=text_content,
            elements=elements,
            metadata=metadata,
            structure_analysis=structure_analysis,
            extraction_errors=errors,
            processing_time=0.0,  # Will be set by caller
            success=len(errors) == 0 and len(text_content) > 0
        )
    
    def _analyze_word(self, document_path: Path) -> AnalysisResult:
        """Analyze Word document"""
        if not DOCX_AVAILABLE:
            raise ImportError("Word processing library not available. Install python-docx.")
        
        text_content = ""
        elements = []
        metadata = DocumentMetadata()
        errors = []
        
        try:
            # Open document
            doc = DocxDocument(str(document_path))
            
            # Extract metadata
            metadata = self._extract_word_metadata(doc)
            
            # Extract text content and structure
            paragraph_texts = []
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:
                    paragraph_texts.append(para_text)
                    
                    # Determine paragraph type
                    para_type = self._classify_paragraph(paragraph)
                    
                    elements.append(DocumentElement(
                        type=para_type,
                        content=para_text,
                        position=para_num,
                        metadata={
                            'paragraph_number': para_num + 1,
                            'style': paragraph.style.name if paragraph.style else 'Normal'
                        }
                    ))
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                if table_text:
                    elements.append(DocumentElement(
                        type=DocumentStructure.TABLE,
                        content=table_text,
                        position=len(elements),
                        metadata={
                            'table_number': table_num + 1,
                            'rows': len(table.rows),
                            'columns': len(table.columns) if table.rows else 0
                        }
                    ))
                    paragraph_texts.append(table_text)
            
            text_content = "\n\n".join(paragraph_texts)
            
            # Analyze document structure
            structure_analysis = self._analyze_text_structure(text_content)
            
            # Update metadata with text statistics
            metadata.word_count = len(text_content.split())
            metadata.character_count = len(text_content)
            
        except Exception as e:
            errors.append(f"Word document analysis error: {str(e)}")
        
        return AnalysisResult(
            document_type=DocumentType.DOCX,
            text_content=text_content,
            elements=elements,
            metadata=metadata,
            structure_analysis=structure_analysis,
            extraction_errors=errors,
            processing_time=0.0,  # Will be set by caller
            success=len(errors) == 0 and len(text_content) > 0
        )
    
    def _extract_pdf_metadata_pdfplumber(self, pdf_metadata: Dict) -> DocumentMetadata:
        """Extract metadata from pdfplumber PDF"""
        return DocumentMetadata(
            title=pdf_metadata.get('Title'),
            author=pdf_metadata.get('Author'),
            subject=pdf_metadata.get('Subject'),
            creator=pdf_metadata.get('Creator'),
            creation_date=self._parse_pdf_date(pdf_metadata.get('CreationDate')),
            modification_date=self._parse_pdf_date(pdf_metadata.get('ModDate'))
        )
    
    def _extract_pdf_metadata_pypdf2(self, pdf_metadata: Any) -> DocumentMetadata:
        """Extract metadata from PyPDF2"""
        return DocumentMetadata(
            title=pdf_metadata.get('/Title'),
            author=pdf_metadata.get('/Author'),
            subject=pdf_metadata.get('/Subject'),
            creator=pdf_metadata.get('/Creator'),
            creation_date=self._parse_pdf_date(pdf_metadata.get('/CreationDate')),
            modification_date=self._parse_pdf_date(pdf_metadata.get('/ModDate'))
        )
    
    def _extract_word_metadata(self, doc: Any) -> DocumentMetadata:
        """Extract metadata from Word document"""
        core_props = doc.core_properties
        
        return DocumentMetadata(
            title=core_props.title,
            author=core_props.author,
            subject=core_props.subject,
            creator=core_props.author,
            creation_date=core_props.created,
            modification_date=core_props.modified
        )
    
    def _parse_pdf_date(self, date_str: Optional[str]) -> Optional[datetime]:
        """Parse PDF date string to datetime"""
        if not date_str:
            return None
        
        try:
            # PDF date format: D:YYYYMMDDHHmmSSOHH'mm'
            if date_str.startswith('D:'):
                date_str = date_str[2:]
            
            # Extract basic date components
            if len(date_str) >= 8:
                year = int(date_str[:4])
                month = int(date_str[4:6])
                day = int(date_str[6:8])
                
                hour = int(date_str[8:10]) if len(date_str) >= 10 else 0
                minute = int(date_str[10:12]) if len(date_str) >= 12 else 0
                second = int(date_str[12:14]) if len(date_str) >= 14 else 0
                
                return datetime(year, month, day, hour, minute, second)
        except (ValueError, IndexError):
            pass
        
        return None
    
    def _classify_paragraph(self, paragraph: Any) -> DocumentStructure:
        """Classify paragraph type based on content and style"""
        text = paragraph.text.strip()
        
        if not text:
            return DocumentStructure.PARAGRAPH
        
        # Check if it's a header (common patterns)
        if (len(text) < 100 and 
            (text.isupper() or 
             any(pattern in text.lower() for pattern in ['declaration', 'conformity', 'section', 'chapter']) or
             paragraph.style.name.startswith('Heading'))):
            return DocumentStructure.HEADER
        
        # Check if it's a list item
        if re.match(r'^\s*[-•\*\d+\.\)]\s+', text):
            return DocumentStructure.LIST
        
        return DocumentStructure.PARAGRAPH
    
    def _extract_table_text(self, table: Any) -> str:
        """Extract text from Word table"""
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(" | ".join(row_data))
        
        return "\n".join(table_data)
    
    def _analyze_text_structure(self, text: str) -> Dict[str, Any]:
        """Analyze text structure and patterns"""
        lines = text.split('\n')
        
        structure_analysis = {
            'total_lines': len(lines),
            'non_empty_lines': len([line for line in lines if line.strip()]),
            'paragraphs': len([line for line in lines if len(line.strip()) > 50]),
            'potential_headers': 0,
            'potential_lists': 0,
            'has_tables': False,
            'sections': [],
            'key_phrases': []
        }
        
        # Analyze line types
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for headers (short lines, often capitalized)
            if len(line) < 100 and (line.isupper() or line.istitle()):
                structure_analysis['potential_headers'] += 1
            
            # Check for list items
            if re.match(r'^\s*[-•\*\d+\.\)]\s+', line):
                structure_analysis['potential_lists'] += 1
            
            # Check for table indicators
            if '|' in line or '\t' in line:
                structure_analysis['has_tables'] = True
        
        # Extract key phrases for EU DoC documents
        eu_doc_patterns = [
            r'declaration\s+of\s+conformity',
            r'ce\s+marking',
            r'medical\s+device',
            r'regulation\s+\(eu\)',
            r'harmonised\s+standard',
            r'conformity\s+assessment',
            r'authorized\s+representative',
            r'notified\s+body'
        ]
        
        for pattern in eu_doc_patterns:
            matches = re.findall(pattern, text.lower())
            if matches:
                structure_analysis['key_phrases'].extend(matches)
        
        return structure_analysis
    
    def _post_process_analysis(self, result: AnalysisResult, document_path: Path, file_size: int) -> AnalysisResult:
        """Post-process analysis results"""
        # Update file size in metadata
        result.metadata.file_size = file_size
        
        # Apply content preprocessing if enabled
        if self.config['content_preprocessing']:
            result.text_content = self._preprocess_text(result.text_content)
        
        # Language detection (simplified)
        if self.config['language_detection']:
            result.metadata.language = self._detect_language(result.text_content)
        
        return result
    
    def _preprocess_text(self, text: str) -> str:
        """Preprocess extracted text"""
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Clean up common artifacts
        text = text.replace('\x0c', '')  # Form feed
        text = text.replace('\x00', '')  # Null characters
        
        return text.strip()
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection based on common words"""
        # Simplified language detection
        english_indicators = ['the', 'and', 'of', 'to', 'a', 'in', 'is', 'it', 'you', 'that']
        german_indicators = ['der', 'die', 'und', 'in', 'den', 'von', 'zu', 'das', 'mit', 'sich']
        
        text_lower = text.lower()
        
        english_count = sum(1 for word in english_indicators if word in text_lower)
        german_count = sum(1 for word in german_indicators if word in text_lower)
        
        if english_count > german_count:
            return 'en'
        elif german_count > 0:
            return 'de'
        else:
            return 'unknown'
    
    def get_analysis_statistics(self) -> Dict[str, Any]:
        """Get analyzer statistics"""
        stats = self.analysis_stats.copy()
        
        if stats['documents_analyzed'] > 0:
            stats['success_rate'] = stats['successful_analyses'] / stats['documents_analyzed']
            stats['average_processing_time'] = stats['total_processing_time'] / stats['documents_analyzed']
        else:
            stats['success_rate'] = 0.0
            stats['average_processing_time'] = 0.0
        
        return stats
    
    def validate_document_compatibility(self, document_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Validate document compatibility without full analysis
        
        Args:
            document_path: Path to document
            
        Returns:
            Validation result dictionary
        """
        document_path = Path(document_path)
        
        validation_result = {
            'is_compatible': False,
            'document_type': DocumentType.UNKNOWN,
            'file_size_mb': 0.0,
            'errors': [],
            'warnings': []
        }
        
        try:
            # Check file exists
            if not document_path.exists():
                validation_result['errors'].append(f"File not found: {document_path}")
                return validation_result
            
            # Check file size
            file_size = document_path.stat().st_size
            file_size_mb = file_size / (1024 * 1024)
            validation_result['file_size_mb'] = file_size_mb
            
            if file_size_mb > self.config['max_file_size_mb']:
                validation_result['errors'].append(f"File too large: {file_size_mb:.1f}MB (max: {self.config['max_file_size_mb']}MB)")
            
            # Check file format
            doc_type = self._detect_document_type(document_path)
            validation_result['document_type'] = doc_type
            
            if doc_type == DocumentType.UNKNOWN:
                validation_result['errors'].append(f"Unsupported file format: {document_path.suffix}")
            elif doc_type == DocumentType.PDF and not PDF_AVAILABLE:
                validation_result['errors'].append("PDF processing libraries not available")
            elif doc_type in [DocumentType.DOCX, DocumentType.DOC] and not DOCX_AVAILABLE:
                validation_result['errors'].append("Word processing library not available")
            
            # Set compatibility
            validation_result['is_compatible'] = len(validation_result['errors']) == 0
            
        except Exception as e:
            validation_result['errors'].append(f"Validation error: {str(e)}")
        
        return validation_result


def create_document_analyzer(config: Optional[Dict[str, Any]] = None) -> DocumentAnalyzer:
    """
    Create and return a DocumentAnalyzer instance
    
    Args:
        config: Optional configuration dictionary
        
    Returns:
        Configured DocumentAnalyzer instance
    """
    return DocumentAnalyzer(config=config)
